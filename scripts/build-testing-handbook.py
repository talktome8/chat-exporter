from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Chat-Exporter-2.0-Testing-and-Monitoring-Handbook.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"
MUTED = "586A7A"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), color)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margin(cell)


def set_run_font(run, size=10.5, bold=False, color=NAVY):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(p, before=0, after=6, line=1.25, align=WD_ALIGN_PARAGRAPH.RIGHT):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    p.alignment = align


def paragraph(container, text="", size=10.5, bold=False, color=NAVY, before=0, after=6, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = container.add_paragraph()
    set_para(p, before, after, 1.25, align)
    r = p.add_run(text)
    set_run_font(r, size, bold, color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    style = doc.styles[f"Heading {level}"]
    p.style = style
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run_font(r, (16, 13, 12)[level - 1], True, (BLUE, BLUE, DARK_BLUE)[level - 1])
    return p


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    for cell, text in zip(t.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.1, WD_ALIGN_PARAGRAPH.RIGHT)
        set_run_font(p.add_run(text), 9.5, True, NAVY)
    for row in rows:
        cells = t.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            set_para(p, 0, 0, 1.15, WD_ALIGN_PARAGRAPH.RIGHT)
            set_run_font(p.add_run(text), 9.2, False, NAVY)
    paragraph(doc, "", size=2, after=3)
    return t


def callout(doc, label, text, color=LIGHT_GRAY):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    set_table_geometry(t, [9360])
    cell = t.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.2, WD_ALIGN_PARAGRAPH.RIGHT)
    set_run_font(p.add_run(f"{label}: "), 10, True, NAVY)
    set_run_font(p.add_run(text), 10, False, NAVY)
    paragraph(doc, "", size=2, after=4)


def page_break(doc):
    doc.add_page_break()


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    set_para(header, 0, 0, 1, WD_ALIGN_PARAGRAPH.LEFT)
    set_run_font(header.add_run("CHAT EXPORTER 2.0 | QA & MONITORING"), 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    set_para(footer, 0, 0, 1, WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(footer.add_run("מסמך עבודה ניתן לעריכה | אין בו מידע משיחות אישיות"), 8.5, False, MUTED)


def add_service_card(doc, service, status):
    heading(doc, f"{service} — {status}", 2)
    table(doc, ["בדיקה", "תוצאה", "ראיה / קישור", "הערות"], [
        ("Widget מופיע פעם אחת", "☐ עבר  ☐ נכשל  ☐ חלקי", "", ""),
        ("Quick export: סדר, טקסט וכפילויות", "☐ עבר  ☐ נכשל  ☐ חלקי", "", ""),
        ("Full scan: התחלה וסוף מאומתים", "☐ עבר  ☐ נכשל  ☐ חלקי", "", ""),
        ("Markdown + TXT + הורדה", "☐ עבר  ☐ נכשל  ☐ חלקי", "", ""),
        ("SPA navigation: אין היעלמות/שכפול", "☐ עבר  ☐ נכשל  ☐ חלקי", "", ""),
    ], [3050, 1900, 2200, 2210])
    paragraph(doc, "דפדפן: ____________________  גירסת תוסף: ____________________  בודק/ת: ____________________", 9.2, False, MUTED, after=9)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in ((1, 16, BLUE, 14, 8), (2, 13, BLUE, 11, 6), (3, 12, DARK_BLUE, 8, 4)):
        s = styles[f"Heading {level}"]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        s._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.15
    add_header_footer(doc)

    # Cover
    paragraph(doc, "CHAT EXPORTER 2.0", 13, True, BLUE, before=70, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "חוברת בדיקות, ניטור ושער שחרור", 27, True, NAVY, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "מדריך עריך לשמירה על ייצוא אמין כאשר ממשקי שירותי ה‑AI משתנים", 13, False, MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    callout(doc, "מטרת המסמך", "לזהות שינוי לפני שמשתמשים נפגעים, להוכיח שהייצוא שלם, ולהוציא עדכון רק כאשר יש צורך מוכח.", LIGHT_BLUE)
    table(doc, ["פריט", "ערך"], [
        ("גרסת מסמך", "1.0"),
        ("תאריך", date.today().isoformat()),
        ("היקף", "Chrome, Edge, Firefox | ChatGPT, Claude, Gemini, Copilot, Perplexity"),
        ("בעלות על אישור שחרור", "בעל המוצר + בודק/ת UAT"),
        ("עקרון פרטיות", "אין שמירת שיחות או כתובות שיחה במערכת הניטור"),
    ], [2500, 6860])
    paragraph(doc, "שימוש: מלאו את השדות, שמרו ראיות שאינן כוללות מידע אישי, ואל תסמנו Complete בלי הוכחת התחלה, סוף, ספירה וסדר.", 10, False, MUTED, before=12, after=0)

    page_break(doc)
    heading(doc, "1. התשובה הקצרה: אוטומציה כן, עדכון עצמי עיוור לא", 1)
    paragraph(doc, "אפשר לבנות מסלול תחזוקה שמזהה שינוי בממשק של שירות AI ומודיע בדיוק מה נשבר. אי אפשר — ולא נכון — לתת לתוסף לשנות selectors, לפרסם לחנות או לגשת לשיחות של משתמשים באופן אוטומטי.")
    table(doc, ["המערכת יכולה לבצע אוטומטית", "דורש תיקון ואישור אנושי"], [
        ("הרצת fixtures, בדיקות יחידה, packaging ואימות ZIP/manifest", "שינוי קוד extractor או selectors אחרי שינוי DOM"),
        ("בדיקת שיחת QA ייעודית מול כל שירות, אם החשבון מחובר", "פרסום Chrome / Edge / Firefox או העלאת גרסה"),
        ("השוואת חוזה DOM מצומצם: קיימת הודעת משתמש, תשובת AI, composer ועוגן widget", "קבלת תנאי שירות, CAPTCHA, login חדש או הרשאות חדשות"),
        ("דוח שינוי, צילום ראיה ופתיחת משימת תחזוקה", "גישה לתוכן של חשבונות משתמשים או לתכתובות אישיות"),
    ], [4680, 4680])
    callout(doc, "כלל בטיחות", "התראה על שינוי היא טריגר לבדיקה; היא אינה הוכחה לתקלה ואינה הרשאה לפרסם עדכון.", "FFF8E8")
    heading(doc, "מה נחשב לשינוי המחייב בדיקה", 2)
    table(doc, ["חומרה", "דוגמה", "תגובה"], [
        ("אדום", "לא נמצא turn, composer, או שהסריקה חוזרת בלי התקדמות", "עצירת Verified, תיקון adapter והרצת מטריצה מלאה"),
        ("צהוב", "שם מודל או עוגן widget השתנו אך הייצוא תקין", "פתיחת משימה; בדיקה ידנית בתוך 72 שעות"),
        ("ירוק", "שינוי CSS שלא משפיע על חוזה extraction", "תיעוד בלבד ובדיקה מחזורית הבאה"),
    ], [1300, 4300, 3760])

    page_break(doc)
    heading(doc, "2. מסלול הניטור המומלץ", 1)
    paragraph(doc, "זהו עיצוב ההפעלה המומלץ לשלב הבא. הוא מונע את הצורך לאתר ידנית, אך משאיר את ההחלטה על קוד ופרסום תחת שליטה.")
    table(doc, ["שלב", "תדירות", "פלט נדרש", "החלטה"], [
        ("A. בדיקות fixture מקומיות", "כל שינוי קוד", "סדר הודעות, כפילויות אמיתיות, filtering, ZIP, SHA-256", "כישלון חוסם build"),
        ("B. בדיקת חוזה DOM", "שבועי + לאחר דיווח תקלה", "קיימים selectors קנוניים וה-widget פעם אחת", "התראה אם חסר חוזה"),
        ("C. שיחת QA חיה ייעודית", "לפני release", "Quick + Full, קובץ, manifest, צילום ראיה", "Pass / Partial / Fail"),
        ("D. שער שחרור", "לכל גרסה", "כל המטריצה, checksum, UAT חתום", "רק אז שינוי גרסה והעלאה"),
    ], [2000, 1600, 3650, 2110])
    heading(doc, "הגדרת חשבון QA", 2)
    table(doc, ["כלל", "יישום"], [
        ("חשבון נפרד", "חשבון בדיקה ייעודי לכל ספק, ללא שיחות פרטיות וללא מידע רגיש."),
        ("שיחה קנונית", "אותו prompt קצר: עברית/אנגלית, קוד, טבלה, קישור והודעה חוזרת מכוונת."),
        ("ראיות מינימליות", "ספירות, hashes וצילומי UI ללא טקסט פרטי. אין להעלות exports אמיתיים למערכת ניטור."),
        ("התראות", "הודעה לבעל המוצר עם: שירות, חומרה, selector חסר, build אחרון וקישור לדוח."),
    ], [2200, 7160])
    callout(doc, "חשוב", "התקנת build unpacked לצד גרסת חנות באותו פרופיל עלולה להציג שני widgets. בדיקות beta יש לבצע בפרופיל Edge/Chrome נפרד, או להשבית זמנית את גרסת החנות באופן הפיך.", "FCEDEC")

    page_break(doc)
    heading(doc, "3. Checklist קדם-שחרור — כל שירות וכל דפדפן", 1)
    paragraph(doc, "מלאו שורה אחת לכל שירות ודפדפן. כל חמשת השירותים כפופים לאותו שער בדיקות.")
    table(doc, ["שירות", "Chrome", "Edge", "Firefox", "מצב", "בעלות / ראיה"], [
        ("ChatGPT", "☐", "☐", "☐", "☐ Verified ☐ Blocked", ""),
        ("Claude", "☐", "☐", "☐", "☐ Verified ☐ Blocked", ""),
        ("Gemini", "☐", "☐", "☐", "☐ Verified ☐ Blocked", ""),
        ("Copilot", "☐", "☐", "☐", "☐ Verified ☐ Blocked", ""),
        ("Perplexity", "☐", "☐", "☐", "☐ Verified ☐ Blocked", ""),
    ], [1280, 900, 850, 1000, 1800, 3530])
    heading(doc, "כל סמן מחייב את כל הבדיקות הבאות", 2)
    table(doc, ["תחום", "בדיקת קבלה"], [
        ("דיוק", "אותה הודעה שנשלחה פעמיים נשמרת פעמיים; אותו node שנלכד בשני selectors נשמר פעם אחת."),
        ("שלמות", "Full scan מגיע להתחלה ולסוף; אחרת התוצאה Partial עם סיבה, ולא מוצגת כהשלמה."),
        ("מבנה", "עברית/אנגלית, קוד, טבלה, רשימה וקישור נשמרים. סדר messages זהה למסך."),
        ("קבצים", "הסינון מיושם לפני הפיצול; מעל 10MiB נוצר ZIP עם חלקים, manifest, SHA-256 וטווחי הודעות."),
        ("UI", "Widget מוטמע פעם אחת ליד composer או fallback צף; popup הוא fallback מלא; מעבר SPA אינו משכפל."),
        ("פרטיות", "אין רשת, backend או storage של תוכן השיחה. settings בלבד ב-storage.local."),
    ], [1700, 7660])

    page_break(doc)
    heading(doc, "4. רישום תוצאות לפי שירות — חלק א׳", 1)
    add_service_card(doc, "ChatGPT", "Verified candidate")
    add_service_card(doc, "Claude", "Verified candidate")
    add_service_card(doc, "Gemini", "Verified candidate")

    page_break(doc)
    heading(doc, "5. רישום תוצאות לפי שירות — חלק ב׳", 1)
    add_service_card(doc, "Copilot", "Verified candidate")
    add_service_card(doc, "Perplexity", "Verified candidate")

    page_break(doc)
    heading(doc, "6. בדיקת ייצוא ארוך ו-ZIP", 1)
    table(doc, ["בדיקה", "קריטריון קבלה", "תוצאה / ראיה"], [
        ("סף קובץ", "עד 10MiB: Markdown או TXT יחיד. מעל 10MiB: ZIP יחיד.", "☐ עבר ☐ נכשל"),
        ("פיצול", "החיתוך רק בין הודעות. הודעה גדולה יחידה נשארת בחלק עצמאי.", "☐ עבר ☐ נכשל"),
        ("שמות", "conversation-YYYY-MM-DD-part-001-of-004 עם סיומת תואמת.", "☐ עבר ☐ נכשל"),
        ("metadata", "מלא רק בחלק הראשון; continuation header בשאר החלקים.", "☐ עבר ☐ נכשל"),
        ("manifest", "גרסה, שירות, שלמות, counts, טווחים, bytes, SHA-256 והגדרות.", "☐ עבר ☐ נכשל"),
        ("שחזור", "איחוד החלקים מחזיר מספר הודעות, סדר וטקסט זהים לייצוא המקורי.", "☐ עבר ☐ נכשל"),
    ], [1700, 5450, 2210])
    heading(doc, "שער שחרור", 2)
    table(doc, ["שומר", "אישור", "שם ותאריך"], [
        ("Lint, unit, integration, packaging ו-verification עברו", "☐", ""),
        ("מטריצה טכנית על unpacked builds הושלמה", "☐", ""),
        ("UAT של בעל המוצר הושלם ללא חסם", "☐", ""),
        ("checksums של חבילות סופיות נשמרו", "☐", ""),
        ("Verified / Beta באתר ובחנויות תואמים לראיות", "☐", ""),
        ("אישור מפורש להעלות לחנויות", "☐", ""),
    ], [5400, 900, 3060])
    callout(doc, "כלל החלטה", "אם אין הוכחה, אין Complete. אם ספק לא עבר — הוא נשאר Beta או מוסר מהפרסום כ-Verified.", "FCEDEC")

    page_break(doc)
    heading(doc, "7. יומן שינוי ותחזוקה", 1)
    paragraph(doc, "מלאו כאן כל שינוי שזוהה. אין למחוק כישלון: מסמנים את התיקון, ה-build שבדק אותו והראיה שהחזירה את השירות ל-Verified.")
    table(doc, ["תאריך", "שירות", "מה השתנה", "חומרה", "פעולה", "Build / ראיה", "סגירה"], [
        ("", "", "", "☐ אדום ☐ צהוב ☐ ירוק", "", "", ""),
        ("", "", "", "☐ אדום ☐ צהוב ☐ ירוק", "", "", ""),
        ("", "", "", "☐ אדום ☐ צהוב ☐ ירוק", "", "", ""),
        ("", "", "", "☐ אדום ☐ צהוב ☐ ירוק", "", "", ""),
        ("", "", "", "☐ אדום ☐ צהוב ☐ ירוק", "", "", ""),
        ("", "", "", "☐ אדום ☐ צהוב ☐ ירוק", "", "", ""),
    ], [750, 950, 2100, 1550, 1450, 1450, 1110])
    heading(doc, "כלל הפעלה קבוע", 2)
    paragraph(doc, "ניטור שבועי מפחית הפתעות; הוא אינו מחליף שחרור. תוסף שכבר מותקן אצל המשתמשים יקבל קוד חדש רק אחרי שהגרסה המתוקנת נבנית, עוברת את שער השחרור ומאושרת/מופצת דרך החנות המתאימה.")
    paragraph(doc, "חתימת בעל המוצר: ____________________     תאריך: ____________________", 10, True, NAVY, before=18, after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
